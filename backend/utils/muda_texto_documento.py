from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def _reescreve_paragrafo(p, novo_texto):
    # preserva a fonte definida no modelo (senão o texto pode não caber na célula)
    fonte_nome = None
    fonte_tamanho = None
    fonte_negrito = None
    if p.runs:
        r0 = p.runs[0]
        fonte_nome = r0.font.name
        fonte_tamanho = r0.font.size
        fonte_negrito = r0.font.bold
    for run in p.runs:
        run.clear()
    p.clear()
    run = p.add_run(novo_texto)
    if fonte_nome:
        run.font.name = fonte_nome
    if fonte_tamanho:
        run.font.size = fonte_tamanho
    if fonte_negrito is not None:
        run.font.bold = fonte_negrito
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def muda_texto_documento(doc, campo, valor):
    for p in doc.paragraphs:
        if campo in p.text:
            _reescreve_paragrafo(p, p.text.replace(campo, valor))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if campo in p.text:
                        _reescreve_paragrafo(p, p.text.replace(campo, valor))
