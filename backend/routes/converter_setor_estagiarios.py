# from utils.valida_ambiente_inux import valida_ambiente_pdf_linux
from utils.convert_to_pdf import convert_to_pdf
from utils.muda_texto_documento import muda_texto_documento
from utils.formata_datas import data_atual, pega_final_de_semana, pega_quantidade_dias_mes
from flask import Blueprint, request, jsonify, send_file
from conection_mysql import connect_mysql
# from mysql.connector import Error # Não usado diretamente
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
import os
import zipfile
from datetime import datetime, date, time, timedelta
# Imports necessários para pegar_feriados_mes
from dateutil.easter import easter
import holidays
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

# NOME DO BLUEPRINT CORRIGIDO PARA RESOLVER O ERRO DE IMPORTAÇÃO
bp_converte_setor_estagiario_pdf = Blueprint('bp_converte_setor_estagiario_pdf', __name__) 

# --- Funções Auxiliares de Estilo e Dados (Inalteradas) ---

def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_row_background(row, color_hex):
    for cell in row.cells:
        set_cell_background(cell, color_hex)


def pegar_feriados_mes(ano, mes, estado='AM'):
    br_feriados = holidays.Brazil(state=estado)
    pascoa = easter(ano)
    corpus_christi = pascoa + timedelta(days=60)
    br_feriados[corpus_christi] = "Corpus Christi"

    conexao = connect_mysql()
    cursor = conexao.cursor(dictionary=True)
    feriados_municipais_db = []
    try:
        query_sql = "SELECT data, ponto_facultativo FROM feriados_municipais WHERE estado = %s AND YEAR(data) = %s"
        params = (estado, ano)
        cursor.execute(query_sql, params)
        feriados_municipais_db = cursor.fetchall()
    finally:
        if conexao.is_connected():
            cursor.close()
            conexao.close()

    pontos_facultativos = []
    for feriado_row in feriados_municipais_db:
        data_db = feriado_row['data']
        ponto_facultativo = feriado_row.get('ponto_facultativo', 0)
        data_feriado_obj = None
        if data_db is None:
            continue
        if hasattr(data_db, 'date'):
            data_feriado_obj = data_db.date()
        elif isinstance(data_db, date):
            data_feriado_obj = data_db
        else:
            try:
                data_feriado_obj = date.fromisoformat(str(data_db))
            except ValueError:
                continue
        if data_feriado_obj:
            if ponto_facultativo:
                pontos_facultativos.append(data_feriado_obj)
            br_feriados[data_feriado_obj] = "Feriado Municipal"
            
    feriados_mes = [d for d in br_feriados if d.month == mes]
    pontos_facultativos_mes = [d for d in pontos_facultativos if d.month == mes]
    
    return feriados_mes, pontos_facultativos_mes


def formatar_horario_para_hh_mm_v2(valor_horario):
    if not valor_horario: return ''
    if isinstance(valor_horario, time): return valor_horario.strftime('%H:%M')
    if isinstance(valor_horario, timedelta):
        total_seconds = abs(int(valor_horario.total_seconds()))
        hours = (total_seconds // 3600) % 24
        minutes = (total_seconds % 3600) // 60
        return f"{hours:02}:{minutes:02}"
    if isinstance(valor_horario, str):
        try:
            if valor_horario.count(':') == 2: return datetime.strptime(valor_horario, '%H:%M:%S').strftime('%H:%M')
            elif valor_horario.count(':') == 1: return datetime.strptime(valor_horario, '%H:%M').strftime('%H:%M')
            return valor_horario
        except ValueError: return valor_horario
    return str(valor_horario)

def limpa_nome(nome):
    nome_limpo = re.sub(r'[<>:"|?*\\/]', '', nome).strip()
    nome_limpo = nome_limpo.replace(' ', '_')
    return nome_limpo

# --- Rota Principal (converte_setores_estagiarios_pdf) ---

@bp_converte_setor_estagiario_pdf.route('/api/setores/estagiar/pdf', methods=['POST'])
def converte_setores_estagiarios_pdf():
    conexao_principal = None
    try:
        body = request.json or {}
        setores = body.get('setores')
        mes_body = body.get('mes')
        ano_body = body.get('ano') 

        if not setores or not isinstance(setores, list):
            return jsonify({'erro': 'Nenhum setor selecionado ou formato inválido'}), 400

        if isinstance(mes_body, list) and len(mes_body) > 0:
            mes_body = mes_body[0]
        if not mes_body:
            return jsonify({'erro': 'Mês não informado'}), 400
        
        # --- DEFINIÇÃO FINAL DE ANO E MÊS NUMÉRICO (RESOLUÇÃO DO BUG DO ANO) ---
        if not ano_body:
            ano = datetime.now().year
        else:
            try:
                ano = int(ano_body)
            except ValueError:
                return jsonify({'erro': 'Ano inválido. Deve ser um número.'}), 400

        data_ano_mes_atual = data_atual(mes_body)
        mes_por_extenso_geral = data_ano_mes_atual['mes']
        
        month_to_num = {
            'Janeiro': 1, 'Fevereiro': 2, 'Março': 3, 'Abril': 4, 'Maio': 5, 'Junho': 6,
            'Julho': 7, 'Agosto': 8, 'Setembro': 9, 'Outubro': 10, 'Novembro': 11, 'Dezembro': 12
        }
        mes_numerico = month_to_num.get(mes_body, datetime.now().month)
        if isinstance(ano_body, str):
            ano = int(ano_body) 
            
        print(f"DEBUG: MÊS NUMÉRICO FINAL UTILIZADO: {mes_numerico}")
        print(f"DEBUG: ANO FINAL UTILIZADO: {ano}") 
        # ---------------------------------------------

        arquivos_zip_gerados_todos_setores = []

        for setor_nome in setores:
            
            conexao_principal = connect_mysql()
            cursor = conexao_principal.cursor(dictionary=True)

            # CONSULTA CORRETA: Buscando ESTAGIÁRIOS
            query_estagiarios = "SELECT * FROM estagiarios WHERE secretaria_lotacao = %s"
            cursor.execute(query_estagiarios, (setor_nome,))
            estagiarios = cursor.fetchall()

            if not estagiarios:
                if conexao_principal and conexao_principal.is_connected():
                    cursor.close()
                    conexao_principal.close()
                print(f"Nenhum estagiário encontrado no setor {setor_nome}, pulando.")
                continue

            arquivos_pdf_gerados_neste_setor = []
            setor_limpo = limpa_nome(setor_nome)

            # Pré-cálculo de feriados para o período
            estado_para_feriados = 'AM'
            feriados_mes_corrente, pontos_fac_mes_corrente = pegar_feriados_mes(ano, mes_numerico, estado=estado_para_feriados)
            
            ano_proximo_mes_periodo = ano
            mes_numerico_proximo_periodo = mes_numerico + 1
            if mes_numerico_proximo_periodo > 12:
                mes_numerico_proximo_periodo = 1
                ano_proximo_mes_periodo += 1

            feriados_proximo_mes, pontos_fac_proximo_mes = pegar_feriados_mes(
                ano_proximo_mes_periodo, mes_numerico_proximo_periodo, estado=estado_para_feriados)

            todos_feriados_do_periodo = list(set(feriados_mes_corrente + feriados_proximo_mes))
            todos_pontos_facultativos_do_periodo = list(set(pontos_fac_mes_corrente + pontos_fac_proximo_mes))


            for estagiario in estagiarios:
                template_path = 'FREQUÊNCIA ESTAGIÁRIOS - MODELO.docx' 
                doc = Document(template_path)
                
                # CHAMA A FUNÇÃO DE CÁLCULO E PREENCHIMENTO (CORRIGIDA)
                periodo_formatado, dias_periodo = cria_dias_da_celula_estagiario(
                    doc, ano, mes_numerico, estagiario, todos_feriados_do_periodo, todos_pontos_facultativos_do_periodo
                )
                
                troca_de_dados = {
                    "CAMPO SETOR": estagiario.get('secretaria_lotacao', ''),
                    "CAMPO MES": periodo_formatado, 
                    "CAMPO NOME": estagiario.get('nome', ''),
                    "CAMPO PERIODO": periodo_formatado,
                    "CAMPO ANO": str(ano), 
                    "CAMPO HORARIO": str(estagiario.get('horario', '')),
                    "CAMPO ENTRADA": formatar_horario_para_hh_mm_v2(estagiario.get('horario_entrada', '')),
                    "CAMPO SAÍDA": formatar_horario_para_hh_mm_v2(estagiario.get('horario_saida', '')),
                    "CAMPO CARGO": estagiario.get('cargo', ''),
                }
                
                for placeholder, valor in troca_de_dados.items():
                    if placeholder in ["CAMPO PERIODO", "CAMPO MES"]:
                        muda_texto_documento_periodo(doc, placeholder, valor) 
                    else:
                        muda_texto_documento(doc, placeholder, valor)

                nome_limpo = estagiario.get('nome', 'NOME_PADRAO').strip().replace('/', '_')
                caminho_pasta = f"setor/{setor_limpo}/estagiario/{mes_por_extenso_geral}"
                
                os.makedirs(caminho_pasta, exist_ok=True)
                nome_base = f"FREQUENCIA_{nome_limpo.replace(' ', '_')}"
                docx_path = os.path.abspath(os.path.join(caminho_pasta, f"{nome_base}.docx"))
                pdf_path = os.path.abspath(os.path.join(caminho_pasta, f"{nome_base}.pdf"))
                doc.save(docx_path)
                convert_to_pdf(docx_path, caminho_pasta)
                arquivos_pdf_gerados_neste_setor.append(pdf_path)

                try:
                    # Ajustado para estagiario_id, conforme a discussão anterior
                    cursor.execute(
                        "INSERT INTO arquivos_pdf (servidor_id, caminho_pdf) VALUES (%s, %s)",
                        (estagiario['id'], pdf_path)
                    )
                except Exception as db_err:
                    print(f"AVISO DB: Falha ao inserir PDF de estagiário na tabela 'arquivos_pdf'. Verifique a coluna 'servidor_id'. Erro: {db_err}")


            if arquivos_pdf_gerados_neste_setor:
                zip_path_setor = f"setor/{setor_limpo}/frequencias_estagiarios_{setor_limpo}_{mes_por_extenso_geral}.zip"
                with zipfile.ZipFile(zip_path_setor, 'w') as zipf:
                    for pdf in arquivos_pdf_gerados_neste_setor:
                        zipf.write(pdf, os.path.basename(pdf))
                arquivos_zip_gerados_todos_setores.append(zip_path_setor)
                
                cursor.execute(
                    "INSERT INTO arquivos_zip (setor, mes, caminho_zip, tipo) VALUES (%s, %s, %s, %s)",
                    (setor_nome, mes_por_extenso_geral, zip_path_setor, 'estagiarios_setor')
                )
            
            conexao_principal.commit()
            if conexao_principal and conexao_principal.is_connected():
                cursor.close()
                conexao_principal.close()


        if not arquivos_zip_gerados_todos_setores:
            return jsonify({'message': 'Nenhum arquivo ZIP de estagiários foi gerado.'}), 200

        if len(arquivos_zip_gerados_todos_setores) > 1:
            zip_final_path = f"setor/frequencias_multissetores_estagiarios_{mes_body.replace('/','-')}_{ano}.zip"
            with zipfile.ZipFile(zip_final_path, 'w') as zipf:
                for zip_file in arquivos_zip_gerados_todos_setores:
                    zipf.write(zip_file, os.path.basename(zip_file))
            
            conexao_principal = connect_mysql()
            cursor = conexao_principal.cursor(dictionary=True)
            
            cursor.execute(
                "INSERT INTO arquivos_zip (setor, mes, ano, caminho_zip, tipo) VALUES (%s, %s, %s, %s, %s)",
                ('multissetores_estagiarios', mes_por_extenso_geral, str(ano), zip_final_path, 'multissetores_estagiarios_geral')
            )
            conexao_principal.commit()
            if conexao_principal and conexao_principal.is_connected():
                cursor.close()
                conexao_principal.close()
            return send_file(zip_final_path, mimetype='application/zip', as_attachment=True, download_name=os.path.basename(zip_final_path))
        elif arquivos_zip_gerados_todos_setores:
            return send_file(arquivos_zip_gerados_todos_setores[0], mimetype='application/zip', as_attachment=True, download_name=os.path.basename(arquivos_zip_gerados_todos_setores[0]))
        
        return jsonify({'message': 'Processamento concluído, mas nenhum ZIP para enviar.'}), 200

    except Exception as exception:
        print(f"ERRO ROTA SETOR ESTAGIÁRIOS: {str(exception)}")
        if conexao_principal and conexao_principal.is_connected():
            if 'cursor' in locals() and cursor:
                cursor.close()
            conexao_principal.close()
        return jsonify({'erro': f'Erro ao processar setores: {str(exception)}'}), 500


# --- LÓGICA DE DATAS E CÉLULAS DOS ESTAGIÁRIOS ---

def calcula_periodo_estagiario(ano_ref, mes_ref):
    # CORREÇÃO: Função auxiliar que só calcula o período
    data_inicio = None
    data_fim = None

    # Exceção 1: 21/11/2025 à 20/12/2025 (Referência: Novembro/2025)
    if mes_ref == 11 and ano_ref == 2025:
        data_inicio = datetime(2025, 11, 21)
        data_fim = datetime(2025, 12, 20)

    # Exceção 2: 21/12/2025 à 31/12/2025 (Referência: Dezembro/2025)
    elif mes_ref == 12 and ano_ref == 2025:
        data_inicio = datetime(2025, 12, 21)
        data_fim = datetime(2025, 12, 31)

    # Exceção 3: 01/01/2026 à 20/01/2026 (Referência: Janeiro/2026)
    elif mes_ref == 1 and ano_ref == 2026:
        data_inicio = datetime(2026, 1, 1)
        data_fim = datetime(2026, 1, 20)

    # LÓGICA NORMAL (21 do mês anterior ao 20 do mês atual)
    else:
        mes_anterior = mes_ref - 1
        ano_anterior = ano_ref
        if mes_ref == 1:
            mes_anterior = 12
            ano_anterior = ano_ref - 1

        data_inicio = datetime(ano_anterior, mes_anterior, 21)
        data_fim = datetime(ano_ref, mes_ref, 20)

    # Constrói a lista de dias dentro do período calculado
    dias_periodo = []
    data_iter_calc = data_inicio
    while data_iter_calc <= data_fim:
        dias_periodo.append({
            "dia": data_iter_calc.day,
            "mes": data_iter_calc.month,
            "ano": data_iter_calc.year
        })
        data_iter_calc += timedelta(days=1)

    # Retorna o período formatado e a lista de dias
    return f"{data_inicio.day:02d}/{data_inicio.month:02d} a {data_fim.day:02d}/{data_fim.month:02d}", dias_periodo


def cria_dias_da_celula_estagiario(doc, ano, mes_numerico, estagiario, feriados, pontos_facultativos):
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    
    # CHAMA A FUNÇÃO DE CÁLCULO DE PERÍODO (CORREÇÃO DE ARGUMENTOS)
    periodo_formatado, dias_periodo = calcula_periodo_estagiario(ano, mes_numerico)

    linha_inicial = 8
    table = doc.tables[0]
    
    # formata estilo das linhas padrão
    for row in table.rows:
        row.height = Cm(0.55)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(7)
                    run.font.bold = False
                    
    # remove ou adiciona linhas para se adequar ao período
    total_linhas_dados_template = len(table.rows) - linha_inicial
    dias_no_periodo_atual = len(dias_periodo)

    if total_linhas_dados_template > dias_no_periodo_atual:
        linhas_para_remover = total_linhas_dados_template - dias_no_periodo_atual
        for _ in range(linhas_para_remover):
            ultima_linha = table.rows[-1]
            tr_element = ultima_linha._tr
            tbl_element = table._tbl
            tbl_element.remove(tr_element)
    elif total_linhas_dados_template < dias_no_periodo_atual:
        linhas_para_adicionar = dias_no_periodo_atual - total_linhas_dados_template
        for _ in range(linhas_para_adicionar):
            new_row = table.add_row()
            new_row.height = Cm(0.55)
            new_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            for cell in new_row.cells:
                p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # preenche os dias
    for i, dia_info in enumerate(dias_periodo):
        dia = dia_info["dia"]
        mes_iter = dia_info["mes"]
        ano_dia = dia_info["ano"]

        data_iteracao_atual = date(ano_dia, mes_iter, dia)
        dia_semana = pega_final_de_semana(ano_dia, mes_iter, dia)

        row = table.rows[linha_inicial + i]

        for cell in row.cells:
            cell.text = ""
            for paragraph in cell.paragraphs:
                paragraph.clear()

        dia_cell = row.cells[0]
        dia_paragraph = dia_cell.paragraphs[0]
        dia_run = dia_paragraph.add_run(str(dia))
        dia_run.font.name = "Calibri"
        dia_run.font.size = Pt(8)
        dia_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        is_ponto_facultativo = data_iteracao_atual in pontos_facultativos
        is_feriado = data_iteracao_atual in feriados
        is_recesso = False

        if estagiario.get('recessoinicio') and estagiario.get('recessofinal'):
            recesso_inicio = estagiario['recessoinicio'].date() if hasattr(estagiario['recessoinicio'], 'date') else estagiario['recessoinicio']
            recesso_final = estagiario['recessofinal'].date() if hasattr(estagiario['recessofinal'], 'date') else estagiario['recessofinal']
            if recesso_inicio <= data_iteracao_atual <= recesso_final:
                is_recesso = True

        texto = None
        if dia_semana == 5:
            texto = "SÁBADO"
        elif dia_semana == 6:
            texto = "DOMINGO"
        elif is_ponto_facultativo and dia_semana not in [5, 6]:
            texto = "PONTO FACULTATIVO"
        elif is_recesso and dia_semana not in [5, 6]:
            texto = "RECESSO"
        elif is_feriado and dia_semana not in [5, 6]:
            texto = "FERIADO"

        if texto:
            set_row_background(row, 'C5E0B4')
            celulas_para_marcar = [2, 5, 9, 13]
            for j in celulas_para_marcar:
                if j < len(row.cells):
                    cell = row.cells[j]
                    for p in cell.paragraphs:
                        p.clear()
                    p_cell = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                    run_cell = p_cell.add_run(texto)
                    run_cell.font.bold = True
                    run_cell.font.name = "Calibri"
                    run_cell.font.size = Pt(6)
                    p_cell.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    return periodo_formatado, dias_periodo


def muda_texto_documento_periodo(doc, campo, valor):
    from utils.muda_texto_documento import _reescreve_paragrafo

    for p in doc.paragraphs:
        if campo in p.text:
            _reescreve_paragrafo(p, p.text.replace(campo, valor))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if campo in p.text:
                        _reescreve_paragrafo(p, p.text.replace(campo, valor))