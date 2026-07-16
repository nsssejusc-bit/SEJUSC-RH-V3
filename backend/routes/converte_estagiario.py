# from utils.valida_ambiente_inux import valida_ambiente_pdf_linux
from utils.convert_to_pdf import convert_to_pdf
from utils.muda_texto_documento import muda_texto_documento
from utils.formata_datas import data_atual, pega_final_de_semana, pega_quantidade_dias_mes
from flask import Blueprint, request, jsonify, send_file
from conection_mysql import connect_mysql
from mysql.connector import Error
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
import os
import zipfile
import datetime
from datetime import time, timedelta, datetime
from datetime import date
from dateutil.easter import easter
import holidays
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re


bp_converte_estagiario_pdf = Blueprint('bp_converte_estagiario_pdf', __name__)


def set_cell_background(cell, color_hex):

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_row_background(row, color_hex):
    for cell in row.cells:
        set_cell_background(cell, color_hex)


def pegar_feriados_mes(ano, mes, estado='AM'):
    br_feriados = holidays.Brazil(state=estado)
    pascoa = easter(ano)
    corpus_christi = pascoa + timedelta(days=60)
    br_feriados[corpus_christi] = "Corpus Christi"

    conexao = connect_mysql()
    cursor = conexao.cursor(dictionary=True)
    feriados_municipais_db = []
    try:
        query_sql = "SELECT data, ponto_facultativo FROM feriados_municipais WHERE estado = %s AND YEAR(data) = %s"
        params = (estado, ano)
        cursor.execute(query_sql, params)
        feriados_municipais_db = cursor.fetchall()
    finally:
        if conexao.is_connected():
            cursor.close()
            conexao.close()

    pontos_facultativos = []
    feriados_normais = []
    for feriado_row in feriados_municipais_db:
        data_db = feriado_row['data']
        ponto_facultativo = feriado_row.get('ponto_facultativo', 0)
        data_feriado_obj = None
        if data_db is None:
            continue
        if hasattr(data_db, 'date'):
            data_feriado_obj = data_db.date()
        elif isinstance(data_db, date):
            data_feriado_obj = data_db
        else:
            try:
                data_feriado_obj = date.fromisoformat(str(data_db))
            except ValueError:
                continue
        if data_feriado_obj:
            if ponto_facultativo:
                pontos_facultativos.append(data_feriado_obj)
            else:
                feriados_normais.append(data_feriado_obj)
            br_feriados[data_feriado_obj] = "Feriado Municipal"
    feriados_mes = [d for d in br_feriados if d.month == mes]
    pontos_facultativos_mes = [d for d in pontos_facultativos if d.month == mes]
    return feriados_mes, pontos_facultativos_mes


def formatar_horario_para_hh_mm_v2(valor_horario):

    if not valor_horario:
        return ''

    if isinstance(valor_horario, time):
        return valor_horario.strftime('%H:%M')

    if isinstance(valor_horario, timedelta):
        total_seconds = int(valor_horario.total_seconds())

        if total_seconds < 0:

            total_seconds = abs(total_seconds)

        hours = (total_seconds // 3600) % 24
        minutes = (total_seconds % 3600) // 60
        return f"{hours:02}:{minutes:02}"

    if isinstance(valor_horario, str):
        try:

            if valor_horario.count(':') == 2:
                dt_obj = datetime.strptime(valor_horario, '%H:%M:%S')
                return dt_obj.strftime('%H:%M')
            # Depois como HH:MM
            elif valor_horario.count(':') == 1:
                dt_obj = datetime.strptime(valor_horario, '%H:%M')
                return dt_obj.strftime('%H:%M')
            else:

                return valor_horario
        except ValueError:

            return valor_horario

    return str(valor_horario)


def limpa_nome(nome):
    # Remove caracteres problemáticos para caminhos de diretório e nomes de arquivo
    nome_limpo = re.sub(r'[<>:"|?*\\/]', '', nome).strip()
    # Substitui espaços por underscores
    nome_limpo = nome_limpo.replace(' ', '_')
    return nome_limpo


@bp_converte_estagiario_pdf.route('/api/estagiario/pdf', methods=['POST'])
def converte_estagiario_pdf():
    try:
        body = request.json or {}
        estagiarios_id = body.get('estagiarios', [])
        print(f"DEBUG: Body recebido: {body}")
        print(f"DEBUG: Estagiários recebidos: {estagiarios_id}")

        if not estagiarios_id:
            return jsonify({'erro': 'Nenhum estagiário selecionado'}), 400

        try:
            ids = [int(id) for id in estagiarios_id]
        except ValueError:
            return jsonify({'erro': 'IDs inválidos'}), 400

        mes_body = body.get('mes')
        print(f"DEBUG: Mês recebido: {mes_body}")
        
        # === CAPTURA E CONVERSÃO DO ANO (PASSO 1) ===
        ano_body = body.get('ano') 
        print(f"DEBUG: Ano recebido (do front-end): {ano_body}")
        
        if not mes_body:
            return jsonify({'erro': 'Mês não informado'}), 400
            
        if not ano_body:
            ano = datetime.now().year
            print(f"AVISO: Ano não informado. Usando ano atual: {ano}")
        else:
            try:
                ano = int(ano_body) # Converte para inteiro (Ex: 2026)
            except ValueError:
                return jsonify({'erro': 'Ano inválido. Deve ser um número.'}), 400
        # ====================================================

        # Tratamento para o caso onde mes_body pode ser uma lista
        if isinstance(mes_body, list) and len(mes_body) > 0:
            mes_body = mes_body[0]
            print(f"DEBUG: Mês extraído da lista: {mes_body}")

        data_ano_mes_atual = data_atual(mes_body)
        mes_por_extenso = data_ano_mes_atual['mes']
        mes_numerico = data_ano_mes_atual['mes_numerico']
        
        # === SOLUÇÃO DEFINITIVA: GARANTIR QUE ANO É O VALOR DO FRONT-END (PASSO 2) ===
        # Reafirma a variável 'ano' com o valor INT capturado de 'ano_body' (Ex: 2026)
        # Se 'data_atual' (que não vimos) estava retornando o ano atual (2025), esta linha corrige.
        if isinstance(ano_body, str):
            ano = int(ano_body)
        # =========================================================================
        
        # Confirma o ano no debug (Este deve ser 2026)
        print(f"DEBUG: ANO FINAL UTILIZADO: {ano}") 

        conexao = connect_mysql()
        cursor = conexao.cursor(dictionary=True)

        placeholders = ','.join(['%s'] * len(ids))
        query = f"SELECT * FROM estagiarios WHERE id IN ({placeholders})"
        cursor.execute(query, ids)
        estagiarios = cursor.fetchall()

        if not estagiarios:
            conexao.close()
            return jsonify({'erro': 'Nenhum estagiário encontrado'}), 404

        arquivos_gerados = []

        estado_para_feriados = 'AM'

        feriados_mes_corrente, pontos_fac_mes_corrente = pegar_feriados_mes(
            ano, mes_numerico, estado=estado_para_feriados)

        # Calcula o próximo mês/ano para buscar feriados que caem no período 21-20
        ano_proximo_mes_periodo = ano
        mes_numerico_proximo_periodo = mes_numerico + 1
        if mes_numerico_proximo_periodo > 12:
            mes_numerico_proximo_periodo = 1
            ano_proximo_mes_periodo += 1

        feriados_proximo_mes, pontos_fac_proximo_mes = pegar_feriados_mes(
            ano_proximo_mes_periodo, mes_numerico_proximo_periodo, estado=estado_para_feriados)

        # Combina feriados dos dois meses para cobrir todo o período 21-20 (ou exceção)
        todos_feriados_do_periodo = list(
            set(feriados_mes_corrente + feriados_proximo_mes))
        todos_pontos_facultativos_do_periodo = list(
            set(pontos_fac_mes_corrente + pontos_fac_proximo_mes))

        for estagiario in estagiarios:
            template_path = 'FREQUÊNCIA ESTAGIÁRIOS - MODELO.docx'
            doc = Document(template_path)

            # CHAMA A FUNÇÃO E CAPTURA O PERÍODO FORMATADO AJUSTADO
            # O ano (2026) e mes_numerico (1) são passados aqui.
            periodo_formatado = cria_dias_da_celula(
                doc, ano, mes_numerico, estagiario, todos_feriados_do_periodo, todos_pontos_facultativos_do_periodo)

            print(f"DEBUG: Período formatado (com exceções): {periodo_formatado}")
            print(f"DEBUG: Mês por extenso: {mes_por_extenso}")
            print(f"DEBUG: Mês numérico: {mes_numerico}")

            troca_de_dados = {
                "CAMPO SETOR": estagiario.get('secretaria_lotacao', estagiario.get('setor', '')),
                "CAMPO MES": periodo_formatado,  # Usa o valor retornado
                "CAMPO NOME": estagiario['nome'],
                "CAMPO PERIODO": periodo_formatado,  # Usa o valor retornado
                "CAMPO ANO": str(ano), # Garante que o ano correto (2026) vá para o documento
                "CAMPO HORARIO": str(estagiario.get('horario')),
                "CAMPO ENTRADA": formatar_horario_para_hh_mm_v2(estagiario.get('horario_entrada')),
                "CAMPO SAÍDA": formatar_horario_para_hh_mm_v2(estagiario.get('horario_saida')),
                "CAMPO CARGO": str(estagiario.get('cargo')),
            }

            print(f"DEBUG: Dados para substituição:")
            for key, value in troca_de_dados.items():
                print(f"  {key}: {value}")
            print(f"DEBUG: Verificando - CAMPO MES: '{troca_de_dados['CAMPO MES']}'")
            print(f"DEBUG: Verificando - CAMPO PERIODO: '{troca_de_dados['CAMPO PERIODO']}'")

            for placeholder, valor in troca_de_dados.items():
                if placeholder in ["CAMPO PERIODO", "CAMPO MES"]:
                    print(
                        f"DEBUG: Chamando função especial para {placeholder} com valor '{valor}'")
                    muda_texto_documento_periodo(doc, placeholder, valor)
                else:
                    muda_texto_documento(doc, placeholder, valor)

            nome_limpo = estagiario['nome'].strip()
            setor_limpo = limpa_nome(
                estagiario.get('secretaria_lotacao', estagiario.get('setor', 'SETOR_PADRAO')))
            caminho_pasta = f"setor/{setor_limpo}/estagiario/{mes_por_extenso}/{nome_limpo}"
            os.makedirs(caminho_pasta, exist_ok=True)

            nome_base = f"{nome_limpo.replace(' ', '_')}_FREQUENCIA"
            docx_path = os.path.abspath(os.path.join(
                caminho_pasta, f"{nome_base}.docx"))
            pdf_path = os.path.abspath(os.path.join(
                caminho_pasta, f"{nome_base}.pdf"))

            # Salva o DOCX e CONVERTE para PDF (NÃO MODIFICADO)
            doc.save(docx_path)
            convert_to_pdf(docx_path, caminho_pasta)

            arquivos_gerados.append(pdf_path)

        zip_path = f"setor/frequencias_{mes_por_extenso}.zip"
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for pdf in arquivos_gerados:
                zipf.write(pdf, os.path.basename(pdf))

        cursor.execute(
            "INSERT INTO arquivos_zip (mes, caminho_zip, tipo) VALUES (%s, %s,%s)",
            (mes_por_extenso, zip_path, 'estagiario')
        )

        conexao.commit()
        conexao.close()

        return send_file(
            zip_path,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f'frequencias_estagiarios_{mes_por_extenso}.zip'
        )

    except Exception as exception:
        if 'conexao' in locals():
            conexao.close()
        return jsonify({'erro': f'Erro: {str(exception)}'}), 500

# -----------------------------------------------------------------------------------------------------------------

def cria_dias_da_celula(doc, ano, mes_numerico, estagiario, feriados, pontos_facultativos):
    from datetime import datetime, timedelta, date
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_ROW_HEIGHT_RULE

    def calcula_periodo_estagiario(ano_ref, mes_ref):
        data_inicio = None
        data_fim = None
        periodo_aplicado = "NORMAL"

        print(f"DEBUG_CALC: Calculando período para Mês={mes_ref}, Ano={ano_ref}")

        # Exceção 2: 21/12/2025 à 31/12/2025 (Ref: Dezembro/2025)
        if mes_ref == 12 and ano_ref == 2025:
            data_inicio = datetime(2025, 12, 21)
            data_fim = datetime(2025, 12, 31)
            periodo_aplicado = "EXCECAO_21DEZ_31DEZ"
            print("DEBUG_CALC: Aplicando Exceção: 21/12/2025 a 31/12/2025")

        # Exceção 3: 01/01/2026 à 20/01/2026 (Ref: Janeiro/2026)
        elif mes_ref == 1 and ano_ref == 2026:
            data_inicio = datetime(2026, 1, 1)
            data_fim = datetime(2026, 1, 20)
            periodo_aplicado = "EXCECAO_01JAN_20JAN"
            print("DEBUG_CALC: Aplicando Exceção: 01/01/2026 a 20/01/2026")

        # LÓGICA NORMAL (21 do mês anterior ao 20 do mês atual)
        else:
            mes_anterior = mes_ref - 1
            ano_anterior = ano_ref
            if mes_ref == 1:
                mes_anterior = 12
                ano_anterior = ano_ref - 1

            data_inicio = datetime(ano_anterior, mes_anterior, 21)
            data_fim = datetime(ano_ref, mes_ref, 20)
            periodo_aplicado = "NORMAL"
            print(f"DEBUG_CALC: Aplicando Normal: {data_inicio.date()} a {data_fim.date()}")

        # Constrói a lista de dias dentro do período calculado
        dias_periodo = []
        data_iter_calc = data_inicio
        while data_iter_calc <= data_fim:
            dias_periodo.append({
                "dia": data_iter_calc.day,
                "mes": data_iter_calc.month,
                "ano": data_iter_calc.year
            })
            data_iter_calc += timedelta(days=1)

        return dias_periodo, data_inicio, data_fim

    linha_inicial = 7
    table = doc.tables[0]

    # formata estilo das linhas padrão
    for row in table.rows:
        row.height = Cm(0.55)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(7)
                    run.font.bold = False

    # chama cálculo do período (vai respeitar exceções acima)
    dias_periodo, data_inicio_periodo, data_fim_periodo = calcula_periodo_estagiario(ano, mes_numerico)

    for i, dia_info in enumerate(dias_periodo):
        dia = dia_info["dia"]
        mes_iter = dia_info["mes"]
        ano_dia = dia_info["ano"]

        data_iteracao_atual = date(ano_dia, mes_iter, dia)
        dia_semana = pega_final_de_semana(ano_dia, mes_iter, dia)

        row = table.rows[linha_inicial + i]

        # limpa conteúdo das células da linha
        for cell in row.cells:
            cell.text = ""
            for paragraph in cell.paragraphs:
                paragraph.clear()

        # preenche número do dia na primeira célula
        dia_cell = row.cells[0]
        dia_paragraph = dia_cell.paragraphs[0]
        dia_run = dia_paragraph.add_run(str(dia))
        dia_run.font.name = "Calibri"
        dia_run.font.size = Pt(8)
        dia_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # checa feriados/ponto facultativo/recesso
        is_ponto_facultativo = data_iteracao_atual in pontos_facultativos
        is_feriado = data_iteracao_atual in feriados
        is_recesso = False

        if estagiario.get('recessoinicio') and estagiario.get('recessofinal'):
            recesso_inicio = estagiario['recessoinicio'].date() if hasattr(estagiario['recessoinicio'], 'date') else estagiario['recessoinicio']
            recesso_final = estagiario['recessofinal'].date() if hasattr(estagiario['recessofinal'], 'date') else estagiario['recessofinal']
            if recesso_inicio <= data_iteracao_atual <= recesso_final:
                is_recesso = True

        texto = None
        if dia_semana == 5:
            texto = "SÁBADO"
        elif dia_semana == 6:
            texto = "DOMINGO"
        elif is_ponto_facultativo and dia_semana not in [5, 6]:
            texto = "PONTO FACULTATIVO"
        elif is_recesso and dia_semana not in [5, 6]:
            texto = "RECESSO"
        elif is_feriado and dia_semana not in [5, 6]:
            texto = "FERIADO"

        if texto:
            set_row_background(row, 'C5E0B4')
            celulas_para_marcar = [2, 5, 9, 13]
            for j in celulas_para_marcar:
                if j < len(row.cells):
                    cell = row.cells[j]
                    for p in cell.paragraphs:
                        p.clear()
                    p_cell = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                    run_cell = p_cell.add_run(texto)
                    run_cell.font.bold = True
                    run_cell.font.name = "Calibri"
                    run_cell.font.size = Pt(6)
                    p_cell.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # remove linhas extras do template se houver (mantém apenas o necessário)
    total_linhas_dados_template = len(table.rows) - linha_inicial
    dias_no_periodo_atual = len(dias_periodo)
    if total_linhas_dados_template > dias_no_periodo_atual:
        linhas_para_remover = total_linhas_dados_template - dias_no_periodo_atual
        for _ in range(linhas_para_remover):
            ultima_linha = table.rows[-1]
            tr_element = ultima_linha._tr
            tbl_element = table._tbl
            tbl_element.remove(tr_element)

    # RETORNA o período no formato antigo -> "DD/MM a DD/MM" (SEM ANO)
    return f"{data_inicio_periodo.day:02d}/{data_inicio_periodo.month:02d} a {data_fim_periodo.day:02d}/{data_fim_periodo.month:02d}"


# -----------------------------------------------------------------------------------------------------------------

def muda_texto_documento_periodo(doc, campo, valor):
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt

    print(f"DEBUG: Aplicando fonte pequena para campo: {campo} com valor: {valor}")

    for p in doc.paragraphs:
        if campo in p.text:
            print(f"DEBUG: Encontrou campo {campo} em parágrafo")
            novo_texto = p.text.replace(campo, valor)
            p.clear()
            run = p.add_run(novo_texto)
            run.font.size = Pt(12)
            run.font.name = "Calibri"
            run.font.bold = False
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            print(f"DEBUG: Aplicou fonte {run.font.size} no parágrafo")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if campo in p.text:
                        print(f"DEBUG: Encontrou campo {campo} em célula de tabela")
                        novo_texto = p.text.replace(campo, valor)
                        p.clear()
                        run = p.add_run(novo_texto)
                        run.font.size = Pt(12)
                        run.font.name = "Calibri"
                        run.font.bold = False
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        print(f"DEBUG: Aplicou fonte {run.font.size} na célula")